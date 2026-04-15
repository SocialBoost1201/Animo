import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_style(run, name=u'游ゴシック', size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for run in h.runs:
        set_font_style(run, name=u'游ゴシック', size=16 if level == 1 else 14 if level == 2 else 12)
    return h

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run()
    set_font_style(run, size=size)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        set_font_style(run, size=11)
    return p

def add_table(doc, data, widths):
    table = doc.add_table(rows=1, cols=len(widths))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, (header, width) in enumerate(zip(data[0], widths)):
        hdr_cells[i].text = header
        hdr_cells[i].width = Inches(width)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_font_style(run, bold=True)
    
    for row_data in data[1:]:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font_style(run)

def add_image(doc, image_path, width_inches=5.5):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        return p
    return None

def create_cast_manual():
    doc = Document()
    add_heading(doc, 'Animo キャスト用操作マニュアル (日本語版)', level=1)
    
    add_paragraph(doc, 'このマニュアルでは、キャストの皆様が Animo CMS を使用して日々の業務を行うための手順を解説します。すべての画面は日本語で表示されます。')
    
    add_heading(doc, '1. アカウントのログインと準備', level=2)
    
    add_heading(doc, '1-1. 新規アカウント登録', level=3)
    add_paragraph(doc, '登録用URL: https://club-animo.jp/cast/register', bold=True)
    
    # Insert Real Registration Screenshot
    add_image(doc, 'docs/manuals/images/cast_register.png', width_inches=3.5)
    add_paragraph(doc, '図：キャスト新規登録画面（実機キャプチャ）', italic=True, size=9)
    
    add_heading(doc, '1-2. ログイン方法', level=3)
    add_paragraph(doc, 'ログインURL: https://club-animo.jp/cast/login', bold=True)
    
    # Insert Real Login Screenshot
    add_image(doc, 'docs/manuals/images/cast_login.png', width_inches=3.5)
    add_paragraph(doc, '図：キャストログイン画面（実機キャプチャ）', italic=True, size=9)
    
    add_heading(doc, '2. ダッシュボードの使い方', level=2)
    add_paragraph(doc, 'ログイン後のメイン画面では、今日やるべきことや最新情報がひと目で分かります。')
    
    # Insert Japanese Dashboard Mockup
    add_image(doc, 'docs/manuals/images/cast_dashboard_jp.png', width_inches=4.0)
    add_paragraph(doc, '図：操作イメージ（ダッシュボード）', italic=True, size=9)
    
    add_heading(doc, '3. 主要機能一覧', level=2)
    table_data = [
        ['機能名', '説明'],
        ['本日の確認', '出勤・退勤の連絡、来店予定の報告を行います。'],
        ['翌週シフト提出', '次週以降の出勤希望日を提出します。'],
        ['今日のブログ', 'お客様向けのブログ記事を作成・投稿します。']
    ]
    add_table(doc, table_data, [1.5, 4.5])
    
    doc.save('docs/manuals/cast_manual.docx')

def create_staff_manual():
    doc = Document()
    add_heading(doc, 'Animo スタッフ用操作マニュアル (日本語版)', level=1)
    
    add_paragraph(doc, 'このマニュアルでは、店長・管理者が Animo CMS を使用して店舗運営を行う手順を解説します。')
    
    add_heading(doc, '1. 管理画面へのアクセス', level=2)
    add_paragraph(doc, '管理者URL: https://club-animo.jp/admin/login', bold=True)
    
    # Insert Real Admin Login Screenshot
    add_image(doc, 'docs/manuals/images/admin_login.png', width_inches=5.0)
    add_paragraph(doc, '図：管理者ログイン画面（実機キャプチャ）', italic=True, size=9)
    
    add_heading(doc, '2. ダッシュボード（店舗状況の把握）', level=2)
    add_paragraph(doc, 'ログイン後、店舗の主要な数字（KPI）を確認できます。')
    
    # Insert Japanese Staff Dashboard Mockup
    add_image(doc, 'docs/manuals/images/staff_dashboard_jp.png', width_inches=5.8)
    add_paragraph(doc, '図：店舗管理システムの全体画面イメージ', italic=True, size=9)
    
    table_data = [
        ['項目', '解説'],
        ['本日の状況', '現在の客数、予約件数、稼働状況を表示します。'],
        ['売上統計', '本日の売上、前日比、今月の累計を表示します。'],
        ['キャスト稼働', '出勤数、稼働中、休憩中の内訳を表示します。']
    ]
    add_table(doc, table_data, [2.0, 4.0])
    
    doc.save('docs/manuals/staff_manual.docx')

if __name__ == '__main__':
    if not os.path.exists('docs/manuals'):
        os.makedirs('docs/manuals')
    create_cast_manual()
    create_staff_manual()
    print('Manuals generated in docs/manuals/')
